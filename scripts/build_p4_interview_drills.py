"""
Build P4_Interview_Drills.docx in Dropbox Interview Prep folder.
Navy/bold-black scan-optimized structure. Colorblind-safe (no red/green).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x00, 0x27, 0x6D)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
OUT_PATH = Path(r"C:\Users\ericg\Dropbox\Job Search\Interview Prep\P4_Interview_Drills.docx")


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.color.rgb = NAVY


def add_label(doc: Document, label: str, text: str) -> None:
    """Add a Q:/L1:/L2: style labeled paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r_label = p.add_run(label + " ")
    r_label.bold = True
    r_label.font.color.rgb = NAVY
    r_label.font.size = Pt(11)
    r_body = p.add_run(text)
    r_body.font.size = Pt(11)
    r_body.font.color.rgb = BLACK


def add_speak_block(doc: Document, label: str, text: str) -> None:
    """Indented speak block with >> marker."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r_label = p.add_run(f"{label}  ")
    r_label.bold = True
    r_label.font.color.rgb = NAVY
    r_label.font.size = Pt(10)
    r_body = p.add_run(text)
    r_body.font.size = Pt(10)
    r_body.font.color.rgb = DARK_GRAY


def add_rule_box(doc: Document, text: str) -> None:
    """Bordered info box for standing rules."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, "E8EEF9")
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = NAVY
    r.bold = True


def add_results_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    for i, label in enumerate(["Probe", "Result"]):
        cell = hdr.cells[i]
        _set_cell_bg(cell, "002769")
        p = cell.paragraphs[0]
        r = p.add_run(label)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
    for row_idx, (probe, result) in enumerate(rows):
        row = tbl.rows[row_idx + 1]
        row.cells[0].paragraphs[0].add_run(probe).font.size = Pt(10)
        row.cells[1].paragraphs[0].add_run(result).font.size = Pt(10)


def build() -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("P4 Interview Drills")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Agentic RCM Pre-Submission Prevention Pipeline")
    r2.font.size = Pt(12)
    r2.font.color.rgb = DARK_GRAY

    doc.add_paragraph()

    add_rule_box(doc,
        "Rule: Every answer starts at L1. Escalate to L2 only on 'tell me more.' L3 only on 'go deeper.'\n"
        "Ownership gate: explain cold + defend the tradeoff + cite what was rejected."
    )
    doc.add_paragraph()

    # ─── ADR-001 ───────────────────────────────────────────────────────────
    add_heading(doc, "ADR-001 — Kafka vs Alternatives")

    add_label(doc, "Q:", "Why Kafka over Kinesis / Spark Streaming / Redpanda?")
    add_label(doc, "L1:", "Kafka is the standard for real-time claims streaming. I needed per-payer ordering and zero-downtime rule updates — both require primitives Kafka has natively.")
    add_label(doc, "L2:", "The partition key = payer_id guarantees all in-flight claims for a given payer land in the same partition and are processed by one consumer thread. That matters because I hot-swap NCCI quarterly editions via a compacted rules.control topic. If two claims from the same payer straddled a rule update on different partitions, one would be scored against stale rules. Kinesis has no compacted-topic primitive. Dagster sensors are polling, not streaming — architecturally dishonest for a pre-submission window story.")
    add_label(doc, "L3:", "In production I'd add a hash salt to the payer key to prevent hot partitions on large payers like UHC (30%+ of Medicare Advantage volume). For MVP with 6 partitions and 6 payers the distribution is clean. Redpanda would shave microseconds off broker latency but P4 is bounded by the LLM API call (~300ms) — Redpanda's edge is irrelevant here.")

    doc.add_paragraph()
    add_label(doc, "Q:", "UHC could be 25–30% of volume. How does that not create a hot partition?")
    add_label(doc, "L1:", "For MVP it's acceptable — 6 payers, 6 partitions, roughly even. Hot partitions are a production-scale concern I've designed for.")
    add_label(doc, "L2:", "The fix is a composite key with a hash salt — partition on payer_id + NPI_suffix or payer_id + date_bucket. This spreads one large payer across multiple partitions while keeping per-payer grouping for rule consistency.")
    add_label(doc, "L3:", "The tradeoff: spreading one payer across partitions breaks strict global ordering per payer. That's acceptable because ordering only matters for rule-version consistency, and the compacted rules.control topic handles that independently. All consumers read the same active edition regardless of partition. So the composite key is safe without breaking the architecture.")

    doc.add_paragraph()
    add_label(doc, "Q:", "Define hot-swap / compacted topic.")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(
        "The rules.control topic has cleanup.policy=compact — Kafka retains only the latest message per key "
        "(like a key-value store). When CMS releases a new NCCI quarterly edition, I publish a single message "
        "with key='ncci_active_version'. Any consumer that starts — or restarts — reads current state immediately, "
        "not the full history. The running consumer polls on a background thread, detects the new version, and "
        "swaps the in-memory PTP/MUE tables atomically. Claims in-flight finish against the old rules; new claims "
        "pick up the new edition. Zero downtime, no restart required."
    )
    r.font.size = Pt(10)
    r.font.color.rgb = DARK_GRAY

    doc.add_paragraph()

    # ─── ADR-002 ───────────────────────────────────────────────────────────
    add_heading(doc, "ADR-002 — Data Strategy")

    add_label(doc, "Q:", "Your data is synthetic. How is this 'real data'?")
    add_label(doc, "L1:", "No public CMS dataset has claim-level denial codes — those live in the 835 remittance and are never published. Realness lives in the policy and distributions, not the rows.")
    add_label(doc, "L2:", "Every denial in P4 traces to a real NCCI edit or Medicare coverage determination — the same rules CMS uses to adjudicate actual Medicare claims. Charge distributions, code frequencies, and NPI samples come from the real 2024 CMS Provider Utilization file. Denial rates are calibrated to real Transparency-in-Coverage PUF issuer rates. The only synthetic atom is composing individual claim rows from these real aggregate distributions.")
    add_label(doc, "L3:", "DE-SynPUF is ICD-9 and statistically perturbed — not defensible as current. The 2023 CMS Synthetic RIF is Synthea-generated, which would erode P4's differentiator from P2 and P3 (both used Synthea intentionally). A real DUA-gated file is the right production answer but weeks of lead time, institutional affiliation required, and the data could never live in a public repo.")

    doc.add_paragraph()
    add_label(doc, "Q:", "How do you know the denial predictions are accurate if you don't have real outcomes?")
    add_label(doc, "L1:", "The validation is against the NCCI rules themselves, not historical claim rows. If the system flags a CO-97 violation on a PTP pair, you can verify it against the actual NCCI quarterly CSV — that's the ground truth.")
    add_label(doc, "L2:", "The noise-injection eval harness proves LLM lift empirically: inject claims with known violations, measure precision/recall/F1 against the deterministic gate. The holdout control arm closes the loop: 10% of claims flow through unmodified, and we compare actual payer adjudications on intervention vs control to measure clean-claim-rate lift.")

    doc.add_paragraph()

    # ─── ADR-003 ───────────────────────────────────────────────────────────
    add_heading(doc, "ADR-003 — Latency and LLM Gate")

    add_label(doc, "Q:", "What's your latency model and why not call the LLM on every claim?")
    add_label(doc, "L1:", "Every claim runs through a deterministic NCCI check in under a millisecond. Only the ambiguous slice — about 15% of volume — ever calls the LLM. That's the cost and latency defense.")
    add_label(doc, "L2:", "The gate produces three routes: PASS (clean, no action), HARD_FAIL (clear violation, no valid modifier bypass), and AMBIGUOUS (modifier present on a modifier_indicator=1 pair — LLM must verify clinical appropriateness). PASS and simple HARD_FAILs never hit the LLM API. Only AMBIGUOUS — and HARD_FAILs on high-dollar claims that need rationale — do. At $0.003/claim for Sonnet, calling the LLM on 100% of volume at 1M claims/month = $3K/month in LLM costs alone. At 15% touch rate that's $450/month.")
    add_label(doc, "L3:", "At production scale I'd replace AMBIGUOUS routing with an XGBoost binary triage trained on historical gate+LLM outcomes. XGBoost handles the volume cheaply; LLM is called only on medium/high-risk claims. That's Phase 3. For v1 the NCCI gate alone reduces LLM calls by 85% and is fully defensible.")

    doc.add_paragraph()

    # ─── FCA Defense ───────────────────────────────────────────────────────
    add_heading(doc, "FCA Defense")

    add_label(doc, "Q:", "You auto-correct claims autonomously. Why isn't that a False Claims Act violation?")
    add_label(doc, "L1:", "Every auto-correct cites the governing NCCI rule. Tool-use grounds the action in the actual policy — it's not the LLM deciding, it's a tool lookup.")
    add_label(doc, "L2:", "FCA turns on falsity + scienter (knowing/reckless disregard). Tool-use attacks falsity: the correction is derived from a tool return value (lookup_ncci_edit(), get_lcd_policy()), not from model judgment. Confidence gating attacks scienter: the system only auto-corrects at ≥92% confidence on charges ≤$500 — everything above that goes to a human queue. Doubt is routed to humans by design.")
    add_label(doc, "L3:", "The immutable audit log seals it: every action records what was changed, what rule was cited, what the confidence was, and a reversibility flag. The kill-switch drops the system to flag-only mode on drift or error-rate breach — that's an active compliance program. A biller editing a portal with no rationale captured is actually higher FCA risk than this system.")

    doc.add_paragraph()

    # ─── Ownership Drill Results ────────────────────────────────────────────
    add_heading(doc, "Ownership Drill Results (Jun 15 2026)", level=2)

    add_results_table(doc, [
        ("Hot partition / payer key", "PASSED — composite key + salt, ordering decoupled from rule hot-swap"),
        ("NCCI routing walkthrough\n20610 + 99213 + modifier 59", "PASSED — MUE → PTP → modifier check → AMBIGUOUS → claims.scored"),
        ("FCA defense", "PASSED — tool-use (falsity) + confidence gate (scienter) + audit log"),
    ])

    doc.add_paragraph()

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
